#!/usr/bin/env python3
"""
TTS Web App - Flask Backend
Supports: ElevenLabs (premium) only - Vietnamese optimized
"""

from flask import Flask, render_template, request, send_file, jsonify
from gtts import gTTS
import os
from datetime import datetime
import uuid
import base64
import requests
from werkzeug.utils import secure_filename

app = Flask(__name__)

# Tạo thư mục lưu audio files và uploads
AUDIO_DIR = 'audio'
UPLOAD_DIR = 'uploads'
os.makedirs(AUDIO_DIR, exist_ok=True)
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Cấu hình upload
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'doc', 'xlsx', 'xls'}

# API Key cố định (ElevenLabs)
ELEVENLABS_API_KEY = "1946e1c982e80e1f9bf4bb06238ed8804a5a78b78f092568b9d01949597eaf14"

# Google Cloud TTS support
def use_google_cloud_tts(text, lang, api_key, voice_name=None):
    """
    Sử dụng Google Cloud Text-to-Speech API với API key

    Args:
        text: Text to convert
        lang: Language code (vi, en, ja, ko)
        api_key: Service account JSON string
        voice_name: Specific voice name (optional)

    Returns:
        bytes: Audio content
        bool: Success status
        str: Error message if failed
    """
    try:
        from google.cloud import texttospeech
        import google.auth
        from google.oauth2 import service_account
        import json

        # Create credentials from API key
        # If API key is a service account JSON
        try:
            creds_dict = json.loads(api_key)
            credentials = service_account.Credentials.from_service_account_info(creds_dict)
        except:
            return None, False, "Invalid API key format. Please use service account JSON."

        # Initialize the client with credentials
        client = texttospeech.TextToSpeechClient(credentials=credentials)

        # Default voice mapping if no specific voice provided
        default_voice_map = {
            'vi': {'language_code': 'vi-VN', 'name': 'vi-VN-Standard-A'},
            'en': {'language_code': 'en-US', 'name': 'en-US-Neural2-F'},
            'ja': {'language_code': 'ja-JP', 'name': 'ja-JP-Neural2-B'},
            'ko': {'language_code': 'ko-KR', 'name': 'ko-KR-Neural2-A'},
        }

        # Parse language code from voice name if provided
        if voice_name and not voice_name.endswith('-standard'):
            # Voice names like: vi-VN-Standard-A, en-US-Neural2-F, etc.
            parts = voice_name.split('-')
            if len(parts) >= 2:
                language_code = f"{parts[0]}-{parts[1]}"
            else:
                language_code = default_voice_map.get(lang, default_voice_map['vi'])['language_code']

            voice_config = {
                'language_code': language_code,
                'name': voice_name
            }
        else:
            # Use default voice for language
            voice_config = default_voice_map.get(lang, default_voice_map['vi'])

        # Set the text input
        synthesis_input = texttospeech.SynthesisInput(text=text)

        # Build the voice request
        voice = texttospeech.VoiceSelectionParams(
            language_code=voice_config['language_code'],
            name=voice_config['name']
        )

        # Select the audio encoding
        audio_config = texttospeech.AudioConfig(
            audio_encoding=texttospeech.AudioEncoding.MP3
        )

        # Perform the text-to-speech request
        response = client.synthesize_speech(
            input=synthesis_input,
            voice=voice,
            audio_config=audio_config
        )

        return response.audio_content, True, None

    except ImportError:
        return None, False, "google-cloud-texttospeech library not installed. Run: pip install google-cloud-texttospeech"
    except Exception as e:
        return None, False, f"Google Cloud TTS Error: {str(e)}"


# ElevenLabs TTS support
def use_elevenlabs_tts(text, voice_id, api_key):
    """
    Sử dụng ElevenLabs API với API key

    Args:
        text: Text to convert
        voice_id: ElevenLabs voice ID
        api_key: ElevenLabs API key

    Returns:
        bytes: Audio content
        bool: Success status
        str: Error message if failed
    """
    try:
        url = f"https://api.elevenlabs.io/v1/text-to-speech/{voice_id}"

        headers = {
            "Accept": "audio/mpeg",
            "Content-Type": "application/json",
            "xi-api-key": api_key
        }

        # Thử các models ElevenLabs (tốt nhất cho tiếng Việt)
        models_to_try = [
            "eleven_turbo_v2_5",       # V2.5 - Nhanh & tốt nhất cho multilingual
            "eleven_multilingual_v2",   # V2 - Multilingual ổn định
            "eleven_turbo_v2"           # V2 - Fallback
        ]

        for model_id in models_to_try:
            data = {
                "text": text,
                "model_id": model_id,
                "voice_settings": {
                    "stability": 0.5,
                    "similarity_boost": 0.8,
                    "style": 0.0,
                    "use_speaker_boost": True
                }
            }

            response = requests.post(url, json=data, headers=headers)

            if response.status_code == 200:
                print(f"✓ ElevenLabs success with model: {model_id}")
                return response.content, True, None
            elif response.status_code != 422:  # 422 = invalid model, thử model khác
                error_msg = f"ElevenLabs API Error {response.status_code}: {response.text}"
                return None, False, error_msg
            # Nếu 422, thử model tiếp theo

        # Tất cả models đều fail
        error_msg = f"ElevenLabs: All models failed. Last error: {response.status_code} - {response.text}"
        return None, False, error_msg

    except Exception as e:
        return None, False, f"ElevenLabs Error: {str(e)}"


# Helper functions cho file upload
def allowed_file(filename):
    """Kiểm tra file extension có hợp lệ không"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_file(filepath):
    """
    Đọc nội dung text từ file

    Hỗ trợ: TXT, PDF, DOCX, XLSX

    Returns:
        str: Nội dung text
        str: Error message nếu có
    """
    try:
        ext = filepath.rsplit('.', 1)[1].lower()

        # TXT file
        if ext == 'txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read(), None

        # PDF file
        elif ext == 'pdf':
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(filepath)
                text = ''
                for page in reader.pages:
                    text += page.extract_text() + '\n'
                return text.strip(), None
            except ImportError:
                return None, "PyPDF2 chưa được cài đặt. Chạy: pip install PyPDF2"
            except Exception as e:
                return None, f"Lỗi đọc PDF: {str(e)}"

        # DOCX file
        elif ext in ['docx', 'doc']:
            try:
                from docx import Document
                doc = Document(filepath)

                # Đọc toàn bộ text từ paragraphs
                full_text = []

                # 1. Đọc paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text)

                # 2. Đọc text trong tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            # Đọc text trong mỗi cell
                            cell_text = cell.text.strip()
                            if cell_text:
                                full_text.append(cell_text)

                # 3. Đọc headers (nếu có)
                for section in doc.sections:
                    header = section.header
                    for para in header.paragraphs:
                        if para.text.strip():
                            full_text.append(para.text)

                    # 4. Đọc footers (nếu có)
                    footer = section.footer
                    for para in footer.paragraphs:
                        if para.text.strip():
                            full_text.append(para.text)

                text = '\n'.join(full_text)
                return text.strip(), None

            except ImportError:
                return None, "python-docx chưa được cài đặt. Chạy: pip install python-docx"
            except Exception as e:
                return None, f"Lỗi đọc DOCX: {str(e)}"

        # XLSX/XLS file
        elif ext in ['xlsx', 'xls']:
            try:
                from openpyxl import load_workbook
                wb = load_workbook(filepath)
                text = ''
                for sheet in wb.worksheets:
                    for row in sheet.iter_rows(values_only=True):
                        row_text = ' '.join([str(cell) if cell is not None else '' for cell in row])
                        if row_text.strip():
                            text += row_text + '\n'
                return text.strip(), None
            except ImportError:
                return None, "openpyxl chưa được cài đặt. Chạy: pip install openpyxl"
            except Exception as e:
                return None, f"Lỗi đọc Excel: {str(e)}"

        else:
            return None, f"Không hỗ trợ định dạng file .{ext}"

    except Exception as e:
        return None, f"Lỗi đọc file: {str(e)}"


@app.route('/')
def index():
    """Render trang chủ"""
    return render_template('index.html')


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """
    API endpoint: Upload file và trích xuất text

    Hỗ trợ: TXT, PDF, DOCX, XLSX (max 16MB)

    Response:
    {
        "success": true,
        "text": "Extracted text content",
        "filename": "uploaded_file.pdf",
        "char_count": 1234
    }
    """
    try:
        # Kiểm tra file có được gửi không
        if 'file' not in request.files:
            return jsonify({
                'success': False,
                'message': 'Không tìm thấy file'
            }), 400

        file = request.files['file']

        # Kiểm tra filename
        if file.filename == '':
            return jsonify({
                'success': False,
                'message': 'Không có file được chọn'
            }), 400

        # Kiểm tra extension
        if not allowed_file(file.filename):
            return jsonify({
                'success': False,
                'message': f'File không hợp lệ. Chỉ hỗ trợ: {", ".join(ALLOWED_EXTENSIONS)}'
            }), 400

        # Save file
        filename = secure_filename(file.filename)
        unique_filename = f"{uuid.uuid4().hex[:8]}_{filename}"
        filepath = os.path.join(UPLOAD_DIR, unique_filename)
        file.save(filepath)

        # Extract text
        text, error = extract_text_from_file(filepath)

        # Xóa file sau khi đọc xong
        try:
            os.remove(filepath)
        except:
            pass

        if error:
            return jsonify({
                'success': False,
                'message': error
            }), 500

        if not text or not text.strip():
            return jsonify({
                'success': False,
                'message': 'File không có nội dung text hoặc không đọc được'
            }), 400

        # Không giới hạn text length - đọc toàn bộ file
        char_count = len(text)
        warning = ''

        # Cảnh báo nếu quá dài (nhưng vẫn hiển thị hết)
        if char_count > 100000:
            warning = ' ⚠️ File rất dài, có thể cần chia nhỏ khi convert'
        elif char_count > 50000:
            warning = ' ⚠️ File dài, recommend chia thành nhiều đoạn'

        return jsonify({
            'success': True,
            'text': text,
            'filename': filename,
            'char_count': char_count,
            'truncated': False,
            'message': f'✓ Đã đọc {char_count:,} ký tự từ file {filename}{warning}'
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Lỗi: {str(e)}'
        }), 500


@app.route('/api/tts', methods=['POST'])
def text_to_speech():
    """
    API endpoint: Convert text to speech

    Supports: ElevenLabs only (Vietnamese optimized)

    Request body:
    {
        "text": "text to convert",
        "voice": "voice ID"
    }

    Response:
    {
        "success": true,
        "audio_url": "/audio/filename.mp3",
        "message": "Success",
        "provider": "ElevenLabs"
    }
    """
    try:
        data = request.get_json()
        text = data.get('text', '')
        voice = data.get('voice', 'elevenlabs-pNInz6obpgDQGcFmaJgB')  # Default: Adam

        if not text.strip():
            return jsonify({
                'success': False,
                'message': 'Text không được để trống'
            }), 400

        # Tăng giới hạn lên 100,000 ký tự
        if len(text) > 100000:
            return jsonify({
                'success': False,
                'message': 'Text quá dài (max 100,000 ký tự). Vui lòng chia nhỏ thành nhiều đoạn.'
            }), 400

        # Cảnh báo nếu text dài (ElevenLabs works best với <5000 chars)
        if len(text) > 5000:
            print(f"⚠️ Warning: Long text ({len(text)} chars). May take longer to process.")

        # Generate unique filename
        filename = f"{uuid.uuid4().hex[:8]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.mp3"
        filepath = os.path.join(AUDIO_DIR, filename)

        # Chỉ dùng ElevenLabs
        voice_id = voice.replace('elevenlabs-', '') if voice.startswith('elevenlabs-') else voice
        audio_content, success, error = use_elevenlabs_tts(text, voice_id, ELEVENLABS_API_KEY)

        if success and audio_content:
            with open(filepath, 'wb') as out:
                out.write(audio_content)
            provider = f'ElevenLabs V3 ({voice_id[:8]}...)'
        else:
            return jsonify({
                'success': False,
                'message': f'Lỗi ElevenLabs: {error}'
            }), 500

        return jsonify({
            'success': True,
            'audio_url': f'/audio/{filename}',
            'message': f'✓ Đã xuất đoạn ghi âm thành công! (Sử dụng {provider})',
            'provider': provider
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Lỗi: {str(e)}'
        }), 500


@app.route('/audio/<filename>')
def serve_audio(filename):
    """Serve audio file"""
    filepath = os.path.join(AUDIO_DIR, filename)

    if not os.path.exists(filepath):
        return jsonify({
            'success': False,
            'message': 'File không tồn tại'
        }), 404

    return send_file(filepath, mimetype='audio/mpeg')


@app.route('/api/cleanup', methods=['POST'])
def cleanup_old_files():
    """Xóa các file audio cũ hơn 1 giờ"""
    try:
        import time
        current_time = time.time()
        deleted_count = 0

        for filename in os.listdir(AUDIO_DIR):
            filepath = os.path.join(AUDIO_DIR, filename)

            # Kiểm tra file cũ hơn 1 giờ (3600 giây)
            if os.path.isfile(filepath):
                file_age = current_time - os.path.getmtime(filepath)
                if file_age > 3600:
                    os.remove(filepath)
                    deleted_count += 1

        return jsonify({
            'success': True,
            'message': f'Đã xóa {deleted_count} file cũ'
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Lỗi: {str(e)}'
        }), 500


if __name__ == '__main__':
    import webbrowser
    import threading

    def open_browser():
        """Tự động mở browser sau 1.5 giây"""
        import time
        time.sleep(1.5)
        webbrowser.open('http://localhost:5000')

    print("=" * 60)
    print("🎙️  TTS Web App đang chạy!")
    print("📍 Tự động mở browser: http://localhost:5000")
    print("💡 Nhấn Ctrl+C để thoát")
    print("=" * 60)

    # Mở browser trong thread riêng
    threading.Thread(target=open_browser, daemon=True).start()

    # Chạy Flask app (tắt debug mode cho production)
    app.run(debug=False, host='0.0.0.0', port=5000, use_reloader=False)
